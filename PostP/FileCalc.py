from io import BytesIO

from PostP import TablesCalc
from tkinter import filedialog
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from PIL import Image


def prepare_data(input_data):
    user_table = []

    f1 = input_data['f1']
    f2 = input_data['f2']
    A = input_data['A']
    L = input_data['L']
    E = input_data['E']
    sigma = input_data['sigma']
    q = input_data['q']

    user_table.append(('f1, F', 'f2, F', 'A, м^2', 'L, м', 'E, Па', 'σ, Па', 'q, Н/м'))

    for i, bar in enumerate(zip(f1, f2, A, L, E, sigma, q)):
        user_table.append(bar)

    calc_tables = TablesCalc.prepare_tables(input_data, 10)
    for calc_table in calc_tables:
        calc_table.insert(0, ("x", "N(x)", "U(x)", "σ(x)", "[σ]"))

    return  user_table, calc_tables


def create_word_report(user_table, calc_tables, image, figure1, figure2, figure3):
    document = Document()

    document.add_heading("Отчет по расчетной стержневой конструкции", level=1)
    document.add_paragraph("")

    document.add_heading("Введенные данные:", level=2)

    table = document.add_table(rows=1, cols=len(
        user_table[0]))
    table.style = 'Table Grid'

    hdr_cells = table.rows[0].cells
    for j, header in enumerate(user_table[0]):
        hdr_cells[j].text = str(header)

    for row in user_table[1:]:
        row_cells = table.add_row().cells
        for j, item in enumerate(row):
            row_cells[j].text = str(item)

    document.add_paragraph("")

    document.add_heading("Построенная конструкция", level=2)
    image = Image.open(image)
    image.save("temp_image.png")
    document.add_picture("temp_image.png", width=Inches(6))
    last_paragraph = document.paragraphs[-1]
    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph("")

    document.add_heading("Расчетные таблицы для конструкции", level=2)

    for table_data in calc_tables:
        table = document.add_table(rows=1, cols=len(table_data[0]))
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells

        for j, header in enumerate(table_data[0]):
            hdr_cells[j].text = str(header)

        for row in table_data[1:]:
            row_cells = table.add_row().cells
            for j, item in enumerate(row):
                row_cells[j].text = str(item)

        document.add_paragraph("")

    document.add_heading('Построенные эпюры', level=2)

    buf1 = BytesIO()
    figure1.savefig(buf1, format='png')
    buf1.seek(0)
    document.add_picture(buf1, width=Inches(6))
    buf1.close()

    buf2 = BytesIO()
    figure2.savefig(buf2, format='png')
    buf2.seek(0)
    document.add_picture(buf2, width=Inches(6))
    buf2.close()

    buf3 = BytesIO()
    figure3.savefig(buf3, format='png')
    buf3.seek(0)
    document.add_picture(buf3, width=Inches(6))
    buf3.close()

    file_path = filedialog.asksaveasfilename(
        defaultextension=".docx",
        filetypes=[("Word Document", "*.docx")],
        title="Сохранить отчет как"
    )

    if file_path:
        document.save(file_path)
